# ======================================================
# IMPORTS
# ======================================================
from pathlib import Path
from io import BytesIO
import re
import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import  WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from textos_especificacoes import texto_topico1, texto_topico2, texto_topico21, texto_topico22, texto_topico23,texto_topico24, texto_topico25, texto_topico26, texto_topico27, texto_topico28
from PIL import Image
import tempfile

# ======================================================
# XML / ESTILO - HELPERS BASE
# ======================================================
def set_cell_background(cell, color_hex: str) -> None:
    """Define cor de fundo de uma célula (hex sem #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _clear_paragraphs(paragraphs):
    for p in paragraphs:
        p.clear()


def _unlink_and_clear_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _clear_paragraphs(section.header.paragraphs)
    _clear_paragraphs(section.footer.paragraphs)


def _link_header_footer_to_previous(section, link: bool = True):
    section.header.is_linked_to_previous = link
    section.footer.is_linked_to_previous = link


def _preparar_imagem_alta_def(image_path: Path):
    """Garante que a imagem tenha 300 DPI para máxima nitidez no Word."""
    if not image_path.exists():
        return None
        
    # Criamos um arquivo temporário para não sujar sua pasta de trabalho
    tmp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    with Image.open(image_path) as img:
        # Converter para RGB se for necessário (evita erros com CMYK ou canais Alpha complexos)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        # Salva com 300 DPI
        img.save(tmp_img.name, "PNG", dpi=(300, 300))
    return tmp_img.name

# ======================================================
# COLUNAS (SEÇÃO)
# ======================================================
def set_section_one_column(section):
    """Volta para 1 coluna na seção."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)

    cols.set(qn("w:num"), "1")
    cols.set(qn("w:sep"), "0")

    # limpa atributos que podem sobrar
    if cols.get(qn("w:space")) is not None:
        cols.attrib.pop(qn("w:space"))
    if cols.get(qn("w:sep")) is not None:
        cols.attrib.pop(qn("w:sep"))


def set_section_two_columns(section, space_twips: int = 720, line_between: bool = False):
    """Configura 2 colunas na seção (auto-flow do Word)."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)

    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:sep"), "1" if line_between else "0")


# ======================================================
# TABELAS - AUTOFIT ROBUSTO
# ======================================================
def enable_autofit_safe(table) -> None:
    """Habilita autofit + tblLayout=autofit (mais robusto no Word)."""
    table.autofit = True
    tblPr = table._tbl.tblPr

    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)

    tblLayout.set(qn("w:type"), "autofit")


def finalize_autofit(table) -> None:
    """Remove larguras fixas (tblW/tcW) e grid, deixando Word ajustar pelo conteúdo."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is not None:
        tblPr.remove(tblW)

    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        tbl.remove(tblGrid)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is not None:
                tcPr.remove(tcW)


# ======================================================
# INSERÇÃO DE TABELAS NO DOCX
# ======================================================
def inserir_tabela_normal_doc(doc, df, largura_total=None):
    """Tabela normal com header azul e linhas alternadas."""
    if hasattr(df, "data"):
        df = df.data

    df = df.reset_index(drop=True)
    rows, cols = df.shape

    table = doc.add_table(rows=rows + 1, cols=cols)
    enable_autofit_safe(table)

    # Header
    for col in range(cols):
        cell = table.cell(0, col)
        cell.text = str(df.columns[col])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "02124A")

        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(255, 255, 255)

    # Body
    for row in range(rows):
        cor = "c6dbef" if row % 2 == 0 else "FFFFFF"
        for col in range(cols):
            cell = table.cell(row + 1, col)
            cell.text = str(df.iloc[row, col])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, cor)

            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for r in p.runs:
                    r.font.size = Pt(10)

    finalize_autofit(table)
    return table


def inserir_tabela_cruzamento_doc(doc, df, largura_total=None):
    """
    Tabela de cruzamento:
    - Header azul
    - Última linha como 'total' (azul e texto branco)
    - Colunas numéricas convertidas para % por linha (quando possível)
    """
    if hasattr(df, "data"):
        df = df.data

    df = df.reset_index(drop=True)
    rows, cols = df.shape

    table = doc.add_table(rows=rows + 1, cols=cols)
    enable_autofit_safe(table)

    # Header
    for col in range(cols):
        cell = table.cell(0, col)
        cell.text = str(df.columns[col])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "02124A")

        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(255, 255, 255)

    total_linha_idx = rows - 1  # última linha

    for row in range(rows):
        is_total = (row == total_linha_idx)
        cor_fundo = "02124A" if is_total else ("c6dbef" if row % 2 == 0 else "FFFFFF")

        # Total da linha (para %)
        total_linha = None
        try:
            total_linha = sum(
                float(str(df.iloc[row, c]).replace(",", "."))
                for c in range(1, cols)
                if str(df.iloc[row, c]).replace(",", ".").isdigit()
            )
        except Exception:
            total_linha = None

        for col in range(cols):
            cell = table.cell(row + 1, col)
            valor = df.iloc[row, col]

            # Converte para % (colunas >= 1)
            if col >= 1 and total_linha and total_linha > 0:
                try:
                    v = float(str(valor).replace(",", "."))
                    texto = f"{(v / total_linha) * 100:.1f}%"
                except Exception:
                    texto = str(valor)
            else:
                texto = str(valor)

            cell.text = texto
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, cor_fundo)

            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if col >= 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
                for r in p.runs:
                    r.font.size = Pt(10)
                    if is_total:
                        r.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)

    finalize_autofit(table)
    return table


def _centralizar_tabela(tabela):
    tbl_pr = tabela._tbl.tblPr
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)


# ======================================================
# MARKDOWN SIMPLES -> DOCX
# ======================================================
def _espacamento(p):
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def _ajustar_titulo(p, tamanho):
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(tamanho)


def _add_runs_markdown(paragraph, text):
    token_pattern = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|~~.+?~~|`.+?`)")
    pos = 0

    for match in token_pattern.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])

        token = match.group()
        if token.startswith("***"):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("~~"):
            run = paragraph.add_run(token[2:-2])
            run.font.strike = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        pos = end

    if pos < len(text):
        paragraph.add_run(text[pos:])


def inserir_markdown_no_doc(doc, markdown_text: str) -> None:
    for linha in markdown_text.split("\n"):
        linha = linha.rstrip()
        if not linha.strip():
            continue

        if re.fullmatch(r"[\s•\-–—]*", linha):
            continue

        # Headers
        m = re.match(r"^(#{1,6})\s+(.*)", linha)
        if m:
            nivel = len(m.group(1))
            texto = m.group(2)

            if nivel == 1:
                p = doc.add_paragraph(texto, style="Heading 1")
                _ajustar_titulo(p, 16)
            elif nivel == 2:
                p = doc.add_paragraph(texto, style="Heading 2")
                _ajustar_titulo(p, 14)
            elif nivel == 3:
                p = doc.add_paragraph(texto, style="Heading 3")
                _ajustar_titulo(p, 13)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(texto)
                run.bold = True
                run.font.size = Pt(12)

            _espacamento(p)
            continue

        # Listas
        m = re.match(r"^[\-\*\+•]\s*(.+)", linha)
        if m:
            texto = m.group(1)
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            _add_runs_markdown(p, texto)
            _espacamento(p)
            continue

        # Parágrafo normal
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _add_runs_markdown(p, linha)
        _espacamento(p)


# ======================================================
# SESSION STATE (MULTIPAGE SAFE)
# ======================================================
def init_session_state_relatorio():
    defaults = {
        "tabelas_doc_intencoes": [],
        "graficos_doc_intencoes": [],
        "tabelas_doc_questoes": [],
        "graficos_doc_questoes": [],
        "tabelas_doc_localidades": [],
        "graficos_doc_localidades": [],
        "tabelas_doc_abertas": [],
        "graficos_doc_abertas": [],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ======================================================
# SEÇÃO PADRÃO DO RELATÓRIO (HEADER/FOOTER)
# ======================================================
def configurar_secao_horizontal(section, cabecalho: str):
    """
    Configuração padrão da seção do relatório:
    - Landscape
    - Margens
    - Cabeçalho com texto e logo
    - Rodapé com linha azul
    """
    section.different_first_page_header_footer = False

    # Orientação landscape
    section.orientation = WD_ORIENT.LANDSCAPE
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width

    # Margens / distâncias
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.3)
    section.header_distance = Inches(0.02)
    section.footer_distance = Inches(0.05)

    largura_util = section.page_width - section.left_margin - section.right_margin

    # Cabeçalho
    header = section.header
    header.is_linked_to_previous = False
    _clear_paragraphs(header.paragraphs)

    tabela_header = header.add_table(rows=1, cols=3, width=largura_util)
    tabela_header.autofit = False
    tabela_header.columns[0].width = Inches(2)
    tabela_header.columns[1].width = largura_util - Inches(4)
    tabela_header.columns[2].width = Inches(2)

    # Texto centro
    cell_text = tabela_header.cell(0, 1)
    cell_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_text = cell_text.paragraphs[0]
    p_text.text = cabecalho
    p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p_text.runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)

    # Logo direita
    cell_logo = tabela_header.cell(0, 2)
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    logo_path = Path.cwd() / "logos" / "2-logo.png"
    if logo_path.exists():
        p_logo.add_run().add_picture(str(logo_path), width=Inches(1.2))

    # Rodapé (linha azul)
    footer = section.footer
    footer.is_linked_to_previous = False
    _clear_paragraphs(footer.paragraphs)

    largura_total_sangria = section.page_width + section.left_margin
    tabela_footer = footer.add_table(rows=1, cols=1, width=largura_total_sangria)
    tabela_footer.autofit = False
    tabela_footer.columns[0].width = largura_total_sangria

    row = tabela_footer.rows[0]
    row.height = Pt(1)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    tbl_pr = tabela_footer._tbl.tblPr

    # puxa para a esquerda
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(-int(section.left_margin.pt * 20)))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    # layout fixo
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    # bordas
    borders = OxmlElement("w:tblBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "50")
    bottom.set(qn("w:color"), "#7DC9FF")

    for side in ["top", "left", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        borders.append(el)

    borders.append(bottom)
    tbl_pr.append(borders)

    return largura_util


def nova_secao_relatorio(doc, cabecalho: str):
    """Cria uma nova seção do relatório já formatada (NOVA PÁGINA)."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.different_first_page_header_footer = False
    _unlink_and_clear_header_footer(section)
    configurar_secao_horizontal(section, cabecalho)
    return section


# ======================================================
# FULL-PAGE (CAPA / PARTIÇÕES) - SEM HEADER/FOOTER
# ======================================================
def _configurar_secao_fullpage_sem_margens(section):
    _unlink_and_clear_header_footer(section)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)


def _usar_paragrafo_para_imagem_fullpage(doc):
    if doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]
        p.clear()
        return p
    p = doc.add_paragraph()
    p.clear()
    return p


def inserir_capa_fullpage(doc, image_path: Path):
    """Capa: seção 0, landscape, sem margens, imagem high-res full-page."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.orientation = WD_ORIENT.LANDSCAPE
    
    # Ajuste de dimensões para Landscape
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width

    _configurar_secao_fullpage_sem_margens(section)

    # Limpeza de parágrafo inicial
    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        p0 = doc.paragraphs[0]
        p0._element.getparent().remove(p0._element)

    p = _usar_paragrafo_para_imagem_fullpage(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Processamento da imagem para alta definição
    img_processada = _preparar_imagem_alta_def(image_path)
    if img_processada:
        # Forçamos a largura e altura para casar com a seção sem distorção
        p.add_run().add_picture(img_processada, width=section.page_width, height=section.page_height)


def inserir_particao_fullpage(doc, image_path: Path):
    """Partição full-page: nova página, alta definição, imagem ocupa tudo."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.different_first_page_header_footer = True
    _unlink_and_clear_header_footer(section)

    # Zerar margens de forma explícita
    section.top_margin = section.bottom_margin = 0
    section.left_margin = section.right_margin = 0
    section.header_distance = section.footer_distance = 0

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    img_processada = _preparar_imagem_alta_def(image_path)
    if img_processada:
        p.add_run().add_picture(img_processada, width=section.page_width, height=section.page_height)

    return section


# ======================================================
# SUBCAPAS (TABELA 2 COLS COM IMAGEM)
# ======================================================
def criar_subcapa_doc(doc, path_esquerda, path_direita, titulo, espacamento_title):
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    # landscape
    section.orientation = WD_ORIENT.LANDSCAPE
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width

    # sem header/footer/margens
    section.different_first_page_header_footer = True
    _unlink_and_clear_header_footer(section)

    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    # tabela 1x2
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    largura_total_emu = section.page_width.emu
    largura_col_dir = int(largura_total_emu * 0.6)
    largura_col_esq = int(largura_total_emu - largura_col_dir)

    largura_col_esq = max(0, largura_col_esq)
    largura_col_dir = max(0, largura_col_dir)

    # força largura total via XML
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(section.page_width.twips)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    table.columns[0].width = largura_col_esq
    table.columns[1].width = largura_col_dir

    # remove bordas
    tbl_borders = OxmlElement("w:tblBorders")
    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{border}")
        b.set(qn("w:val"), "none")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)

    # cell esquerda
    cell_esq = table.cell(0, 0)
    tc_pr = cell_esq._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for edge in ["top", "left", "bottom", "right"]:
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)

    # logo esquerda
    p_logo = cell_esq.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.left_indent = Inches(0.25)
    p_logo.paragraph_format.space_before = Pt(36)
    p_logo.paragraph_format.line_spacing = 1.0

    if path_esquerda.exists():
        p_logo.add_run().add_picture(str(path_esquerda), width=Inches(3.8))

    # título inferior
    p_titulo = cell_esq.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_titulo.paragraph_format.left_indent = Inches(0.4)
    p_titulo.paragraph_format.space_before = Pt(espacamento_title)
    p_titulo.paragraph_format.space_after = Pt(0)
    p_titulo.paragraph_format.line_spacing = 1.0

    run_tit = p_titulo.add_run(str(titulo))
    run_tit.bold = True
    run_tit.font.size = Pt(30)
    font_name = "Arial"
    run_tit.font.name = font_name
    run_tit.font.color.rgb = RGBColor(0, 51, 103)
    rFonts = run_tit._element.rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)

    # cell direita (imagem)
    cell_dir = table.cell(0, 1)
    p_dir = cell_dir.paragraphs[0]
    p_dir.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dir.paragraph_format.space_before = Pt(0)
    p_dir.paragraph_format.space_after = Pt(0)

    if path_direita.exists():
        p_dir.add_run().add_picture(
            str(path_direita),
            width=largura_col_dir,
            height=section.page_height,
        )

    return section

def _logo_path(nome_arquivo: str) -> Path:
    return Path.cwd() / "logos" / nome_arquivo

# Atalhos “páginas”
def inserir_especificacoes_tecnicas(doc):
    criar_subcapa_doc(
        doc,
        _logo_path("2-logo.png"),
        _logo_path("4-especificacoes.jpg"),
        "ESPECIFICAÇÕES TÉCNICAS",
        310
    )

def inserir_pagina_intencoes(doc):
    criar_subcapa_doc(doc, 
                    _logo_path("2-logo.png"),
                    _logo_path("5-intencoes.png"),
                    "INTENÇÕES\nDE VOTO",
                    310
    )

def inserir_pagina_questoes(doc):
    criar_subcapa_doc(doc, 
                        _logo_path("2-logo.png"),
                        _logo_path("6-sociais.jpg"),
                        "CRUZAMENTO POR SEXO, IDADE, ESCOLARIDADE E RENDA",
                        260
    )

def inserir_imagem_doacoes(doc):
    inserir_particao_fullpage(doc, _logo_path("7-doacoes.png"))

def inserir_imagem_final(doc):
    inserir_particao_fullpage(doc, _logo_path("8-final.png"))

def inserir_subcapa(doc, cabecalho):
    criar_subcapa_doc(
        doc,
        _logo_path("2-logo.png"),
        _logo_path("3-lado_direito.png"),
        cabecalho,
        260,
    )


def inserir_pagina_localidades(doc):
    criar_subcapa_doc(
        doc,
        _logo_path("2-logo.png"),
        _logo_path("9-logo_localidades.png"),
        "CRUZAMENTO POR LOCALIDADES",
        310,
    )

# ======================================================
# Sumário automático (campo do Word, baseado em estilos Heading)
# ======================================================
def inserir_sumario_automatico(doc, titulo="SUMÁRIO"):
    # 1. FORÇAR DEFINIÇÃO DOS ESTILOS TOC (Base para o Word)
    for i in range(1, 4):
        style_name = f'TOC {i}'
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        style.font.name = "Arial"
        style.font.size = Pt(12)
        style.font.bold = True 
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(6)

    # 2. TÍTULO COM ESPAÇAMENTO DE 20 PT ABAIXO
    p = doc.add_paragraph(titulo)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(2, 18, 74)
    # Define o recuo de 20 pontos após o título
    p.paragraph_format.space_after = Pt(20) 

    # 3. PARÁGRAFO DO SUMÁRIO (Com Negrito Forçado na Run)
    p_toc = doc.add_paragraph()
    # Aplicamos o estilo TOC 1 ao parágrafo pai para tentar forçar a herança
    try:
        p_toc.style = doc.styles['TOC 1']
    except:
        pass

    run_toc = p_toc.add_run()
    run_toc.bold = True # Força negrito na run que contém o campo
    
    # Inserção do XML do Campo
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')

    # \o "1-3": Níveis | \h: Hyperlinks | \u: Preserva formatação do parágrafo
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run_toc._r.append(fldChar1)
    run_toc._r.append(instrText)
    run_toc._r.append(fldChar2)
    run_toc._r.append(fldChar3)

    # 4. Instrução de atualização
    p_instr = doc.add_paragraph(
        "Nota: Se o sumário aparecer desatualizado, ou não aparecer,  use Ctrl+A e depois F9 para atualizar tudo."
    )
    p_instr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_instr.runs[0].italic = True
    p_instr.runs[0].font.size = Pt(10)
    p_instr.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

# ======================================================
# ESPECIFICAÇÕES TÉCNICAS - 2 COLUNAS (COM VOLTA P/ 1 COL)
# ======================================================
def add_topic(doc, titulo: str, texto: str):
    """Insere o título com estilo Heading 2 para ser capturado pelo Sumário."""
    # Aplicamos o estilo 'Heading 2' explicitamente
    p = doc.add_paragraph(titulo.strip(), style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # Re-aplicamos a formatação visual (porque o estilo Heading do Word costuma ser azul/grande)
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(2, 18, 74) # Azul Marinho da sua logo
    r.font.name = "Arial"

    # Conteúdo do tópico
    for linha in texto.strip().split("\n"):
        if not linha.strip(): continue
        p2 = doc.add_paragraph(linha)
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.first_line_indent = Inches(0.5)
        for run in p2.runs:
            run.font.size = Pt(11)
            run.font.name = "Arial"


def criar_especificacoes_tecnicas(doc, cabecalho_secao=None):
    """
    Cria a(s) página(s) de Especificações Técnicas em 2 colunas,
    com auto-flow (esquerda -> direita -> próxima página).

    Ao final, volta para 1 coluna SEM nova página (CONTINUOUS).
    """
    # seção 2 colunas (nova página)
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    # landscape + margens
    section.orientation = WD_ORIENT.LANDSCAPE
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # header/footer
    if cabecalho_secao:
        configurar_secao_horizontal(section, cabecalho_secao)
    else:
        _unlink_and_clear_header_footer(section)

    # colunas
    set_section_two_columns(section, space_twips=720)

    # Inserção dos tópicos (agora usando Heading 2 internamente)
    add_topic(doc, "1 – OBJETIVO:", texto_topico1())
    add_topic(doc, "2 – METODOLOGIA:", texto_topico2())
    add_topic(doc, "2.1 – UNIVERSO:", texto_topico21())
    add_topic(doc, "2.2 – LOCAL:", texto_topico22())
    add_topic(doc, "2.3 – DIMENSIONAMENTOS AMOSTRAL:", texto_topico23())
    add_topic(doc, "2.4 – CRITÉRIOS DE AMOSTRAGEM:", texto_topico24())
    add_topic(doc, "2.5 – COLETA DE DADOS:", texto_topico25())
    add_topic(doc, "2.6 – CONTROLE DE QUALIDADE:", texto_topico26())
    add_topic(doc, "2.7 – PROCESSAMENTO DOS DADOS:", texto_topico27())
    add_topic(doc, "2.8 – OBSERVAÇÕES:", texto_topico28())

    section_normal = doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_one_column(section_normal)
    _link_header_footer_to_previous(section_normal, True)

    return section


# ======================================================
# GERAÇÃO DO RELATÓRIO
# ======================================================
def gerar_relatorio_docx(cabecalho: str, titulo_subcapa: str) -> BytesIO:
    init_session_state_relatorio()
    doc = Document()

    # Normaliza estilo base
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # 1) CAPA (seção 0)
    inserir_capa_fullpage(doc, _logo_path("1-capa.png"))

    # 2) Seção principal do relatório (após capa)
    sec_relatorio = doc.add_section(WD_SECTION.NEW_PAGE)
    sec_relatorio.different_first_page_header_footer = False
    _unlink_and_clear_header_footer(sec_relatorio)

    largura_util = configurar_secao_horizontal(sec_relatorio, cabecalho)

    # Dimensões para gráficos
    altura_util = sec_relatorio.page_height - sec_relatorio.top_margin - sec_relatorio.bottom_margin
    largura_grafico = largura_util * 0.9
    altura_grafico = altura_util * 0.75

    # ------------------------------
    # Inserção de gráficos
    # ------------------------------
    def inserir_graficos(graficos):
        for info in graficos:
            grafico = info.get("grafico")
            titulo = info.get("titulo", "")
            if grafico is None:
                continue

            p_title = doc.add_paragraph(titulo, style="Heading 2")
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.space_before = Pt(2)
            p_title.paragraph_format.space_after = Pt(18)
            p_title.runs[0].bold = True
            p_title.runs[0].font.size = Pt(12)

            r = p_title.runs[0]
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(2, 18, 74) # Azul Marinho da sua logo
            r.font.name = "Arial"

            img = BytesIO()
            grafico.savefig(img, format="png", bbox_inches="tight")
            img.seek(0)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img, width=largura_grafico, height=altura_grafico)

            doc.add_page_break()

    # ------------------------------
    # Inserção de tabelas
    # ------------------------------
    def inserir_tabelas(tabelas):
        for info in tabelas:
            df = info.get("tabela")
            titulo = info.get("titulo", "")
            interpretacao = info.get("interpretacao", "")

            if df is None:
                continue

            p_title = doc.add_paragraph(titulo, style="Heading 2")
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after = Pt(18)
            p_title.runs[0].bold = True
            p_title.runs[0].font.size = Pt(12)

            r = p_title.runs[0]
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(2, 18, 74) # Azul Marinho da sua logo
            r.font.name = "Arial"

            if re.search(r"(estimulada|cruzamento|obj)", titulo.lower()):
                tabela = inserir_tabela_cruzamento_doc(doc, df)
            else:
                tabela = inserir_tabela_normal_doc(doc, df)

            _centralizar_tabela(tabela)
            doc.add_page_break()

            if interpretacao.strip():
                inserir_markdown_no_doc(doc, interpretacao)
                doc.add_page_break()

    # ======================
    # SUBCAPA
    # ======================
    inserir_subcapa(doc, titulo_subcapa)
    sec_relatorio = doc.add_section(WD_SECTION.NEW_PAGE)
    configurar_secao_horizontal(sec_relatorio, cabecalho)
    inserir_sumario_automatico(doc)

    # ======================
    # ESPECIFICAÇÕES TÉCNICAS
    # ======================
    inserir_especificacoes_tecnicas(doc)
    criar_especificacoes_tecnicas(doc, cabecalho_secao=cabecalho)

    # ======================
    # INTENÇÕES
    # ======================
    inserir_pagina_intencoes(doc)
    nova_secao_relatorio(doc, cabecalho)

    inserir_graficos(st.session_state.graficos_doc_intencoes)
    inserir_tabelas(st.session_state.tabelas_doc_intencoes)

    # ======================
    # ABERTAS
    # ======================
    inserir_tabelas(st.session_state.tabelas_doc_abertas)
    inserir_graficos(st.session_state.graficos_doc_abertas)

    # ======================
    # LOCALIDADES
    # ======================
    inserir_pagina_localidades(doc)
    nova_secao_relatorio(doc, cabecalho)

    inserir_tabelas(st.session_state.tabelas_doc_localidades)
    inserir_graficos(st.session_state.graficos_doc_localidades)

    # ======================
    # QUESTÕES
    # ======================
    inserir_pagina_questoes(doc)
    nova_secao_relatorio(doc, cabecalho)
    
    inserir_tabelas(st.session_state.tabelas_doc_questoes)
    inserir_graficos(st.session_state.graficos_doc_questoes)

    # ======================
    # FINAL
    # ======================
    inserir_imagem_doacoes(doc)
    inserir_imagem_final(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer